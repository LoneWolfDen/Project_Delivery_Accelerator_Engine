"""DOCX parser – extracts text from .docx (Word) files.

Uses ``python-docx`` (open-source, pure-Python) when available.
Falls back to a clear error message if the library is not installed.

Install:
    pip install python-docx    # already included in extras [docs]
"""

from pathlib import Path

from models.document import (
    ContentType,
    DocumentMetadata,
    DocumentSection,
    IngestedDocument,
    SourceType,
)

# Heading styles that Word uses (matches style name prefix, case-insensitive)
_HEADING_STYLE_PREFIX = "heading"


def parse(file_path: Path) -> IngestedDocument:
    """Parse a .docx file into an IngestedDocument.

    Walks document paragraphs; Word Heading styles become section headings,
    body text is grouped under the most-recent heading.

    Args:
        file_path: Path to the .docx file.

    Returns:
        IngestedDocument with heading-based sections.

    Raises:
        ImportError: If python-docx is not installed.
        ValueError: If the file cannot be opened.
    """
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to parse Word documents. "
            "Install it with:  pip install python-docx"
        ) from exc

    try:
        doc = Document(str(file_path))
    except Exception as exc:
        raise ValueError(f"Could not open DOCX '{file_path.name}': {exc}") from exc

    sections: list[DocumentSection] = []
    current_heading = ""
    current_lines: list[str] = []
    section_start = 0
    line_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            line_num += 1
            continue

        style_name = (para.style.name or "").lower()
        is_heading = style_name.startswith(_HEADING_STYLE_PREFIX)

        if is_heading:
            # Flush previous section
            if current_lines or current_heading:
                sections.append(
                    DocumentSection(
                        heading=current_heading,
                        content="\n".join(current_lines).strip(),
                        section_type="body",
                        line_start=section_start,
                        line_end=line_num - 1,
                    )
                )
            current_heading = text
            current_lines = []
            section_start = line_num
        else:
            current_lines.append(text)

        line_num += 1

    # Flush final section
    if current_lines or current_heading:
        sections.append(
            DocumentSection(
                heading=current_heading,
                content="\n".join(current_lines).strip(),
                section_type="body",
                line_start=section_start,
                line_end=line_num - 1,
            )
        )

    # Fallback: document with no heading styles
    if not sections:
        full_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        sections.append(
            DocumentSection(
                heading="",
                content=full_text,
                section_type="body",
                line_start=0,
                line_end=line_num,
            )
        )

    raw_text = "\n\n".join(
        f"{s.heading}\n{s.content}".strip() for s in sections if s.content
    )

    # Also extract tables as plain text
    table_text_parts: list[str] = []
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_text_parts.append("\n".join(rows))

    if table_text_parts:
        table_blob = "\n\n".join(table_text_parts)
        raw_text = raw_text + "\n\n" + table_blob
        sections.append(
            DocumentSection(
                heading="Tables",
                content=table_blob,
                section_type="body",
                line_start=line_num,
                line_end=line_num + len(table_text_parts),
            )
        )

    metadata = DocumentMetadata(
        title=_extract_title(doc, file_path),
        source_type=_detect_source_type(raw_text, file_path),
        word_count=len(raw_text.split()),
        line_count=line_num,
    )

    return IngestedDocument(
        filename=file_path.name,
        file_path=str(file_path),
        content_type=ContentType.TEXT,
        metadata=metadata,
        sections=sections,
        raw_text=raw_text,
    )


def _extract_title(doc, file_path: Path) -> str:
    """Use doc core properties title if set, else filename stem."""
    try:
        title = doc.core_properties.title
        if title and title.strip():
            return title.strip()
    except Exception:
        pass
    return file_path.stem.replace("_", " ").replace("-", " ").title()


def _detect_source_type(raw_text: str, file_path: Path) -> SourceType:
    """Heuristic source-type detection."""
    lower = raw_text.lower()
    name_lower = file_path.stem.lower()

    if any(k in name_lower for k in ("scope", "sow", "statement_of_work")):
        return SourceType.SOW
    if "proposal" in name_lower:
        return SourceType.PROPOSAL
    if any(k in name_lower for k in ("requirement", "nfr", "spec")):
        return SourceType.REQUIREMENTS
    if any(k in lower for k in ("scope of work", "deliverables", "statement of work")):
        return SourceType.SOW
    if any(k in lower for k in ("action items", "next steps", "meeting notes")):
        return SourceType.CALL_NOTES

    return SourceType.ARTEFACT
